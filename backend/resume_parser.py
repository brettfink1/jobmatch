import io
import re
import logging

logger = logging.getLogger(__name__)


def _clean_text(text: str) -> str:
    """Strip excessive whitespace and normalize newlines."""
    # Replace multiple spaces/tabs with single space within lines
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        line = re.sub(r"[ \t]+", " ", line).strip()
        cleaned_lines.append(line)
    # Collapse 3+ consecutive blank lines into 2
    text = "\n".join(cleaned_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber

    all_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text.append(page_text)
    return "\n".join(all_text)


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Parse a resume file (PDF or DOCX) and return the raw extracted text.

    Args:
        file_bytes: The raw bytes of the uploaded file.
        filename: Original filename, used to determine format.

    Returns:
        Cleaned raw text extracted from the resume.

    Raises:
        ValueError: If the file format is not supported or parsing fails.
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        logger.info("Parsing PDF resume: %s", filename)
        try:
            raw_text = _parse_pdf(file_bytes)
        except Exception as exc:
            logger.error("PDF parsing failed: %s", exc)
            raise ValueError(f"Failed to parse PDF: {exc}") from exc
    elif lower_name.endswith(".docx"):
        logger.info("Parsing DOCX resume: %s", filename)
        try:
            raw_text = _parse_docx(file_bytes)
        except Exception as exc:
            logger.error("DOCX parsing failed: %s", exc)
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc
    else:
        raise ValueError(
            f"Unsupported file format: '{filename}'. Please upload a PDF or DOCX file."
        )

    if not raw_text or not raw_text.strip():
        raise ValueError("No text could be extracted from the resume. The file may be image-based or corrupted.")

    return _clean_text(raw_text)
